import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 创建一个新的Word文档
doc = docx.Document()

# 设置标题样式
title = doc.add_heading('基于Transformer的自然语言处理研究', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加摘要
doc.add_heading('摘要', 1)
abstract = doc.add_paragraph(
    '本文简要讨论了基于Transformer架构的自然语言处理模型的最新进展。我们重点关注BERT和GPT等预训练语言模型在各种NLP任务中的应用。'
    '研究表明，这些模型在文本分类、问答系统和机器翻译等任务上取得了显著的性能提升[1]。然而，这些大型模型也带来了计算资源需求增加和环境影响等问题[2]。'
)

# 添加引言
doc.add_heading('1. 引言', 1)
intro = doc.add_paragraph(
    '自然语言处理(NLP)领域在近年来取得了显著进展，主要得益于Transformer架构的提出和预训练语言模型的发展。'
    'Transformer模型通过自注意力机制有效捕捉序列中的长距离依赖关系，为各种NLP任务提供了强大的基础。\n\n'
    'BERT(Bidirectional Encoder Representations from Transformers)是一种基于Transformer的预训练语言模型，'
    '它通过双向上下文学习获取更丰富的语言表示[1]。BERT的出现显著提高了多种NLP任务的性能基准，包括文本分类、命名实体识别和问答系统等。\n\n'
    '随着模型规模的不断扩大，研究人员也开始关注这些大型语言模型的环境影响和计算成本。'
    'Strubell等人的研究指出，训练大型NLP模型会消耗大量能源并产生相应的碳排放[2]。'
)

# 添加结论
doc.add_heading('2. 结论', 1)
conclusion = doc.add_paragraph(
    '基于Transformer的预训练语言模型已经成为NLP领域的主流技术，在多种任务上取得了显著成功。'
    '然而，随着模型规模的增长，计算资源需求和环境影响也成为亟待解决的问题。'
    '未来的研究方向包括开发更高效的模型架构和训练方法，以及探索如何在保持性能的同时减少资源消耗。'
)

# 添加参考文献
doc.add_heading('参考文献', 1)
references = doc.add_paragraph()
references.add_run('[1] Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2018). BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. arXiv:1810.04805.')
references.add_run('\n\n')
references.add_run('[2] Strubell, E., Ganesh, A., & McCallum, A. (2019). Energy and Policy Considerations for Deep Learning in NLP. arXiv:1906.02243.')

# 保存文档
doc.save('arxiv_test_paper.docx')

print("Word文档已创建: arxiv_test_paper.docx")