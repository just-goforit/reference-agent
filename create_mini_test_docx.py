import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 创建一个新的Word文档
doc = docx.Document()

# 设置标题样式
title = doc.add_heading('深度学习在自然语言处理中的应用研究', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加摘要
doc.add_heading('摘要', 1)
abstract = doc.add_paragraph(
    '随着计算能力的提升和大规模数据的可用性增加，深度学习技术在自然语言处理(NLP)领域取得了显著进展。本文简要介绍了深度学习在NLP中的主要应用，'
    '包括文本分类和机器翻译等。我们分析了Transformer架构[1]及其衍生模型如BERT[2]的影响，并讨论了这些模型在NLP任务中的表现。'
    '研究表明，预训练语言模型显著提高了NLP任务的性能基准(Devlin et al., 2019)，但同时也带来了计算资源需求增加等挑战。'
)

# 添加引言
doc.add_heading('1. 引言', 1)
intro = doc.add_paragraph(
    '自然语言处理(NLP)是人工智能的核心分支之一，致力于使计算机能够理解、解释和生成人类语言。近年来，深度学习技术的发展彻底改变了NLP领域的研究方向和应用前景。'
    '与传统的基于规则和统计的方法相比，深度学习模型能够自动学习语言的复杂特征和模式，无需人工特征工程(Goldberg, 2017)[3]。\n\n'
    '深度学习在NLP中的应用始于词嵌入技术的突破。Word2Vec(Mikolov et al., 2013)[4]等方法能够将单词映射到低维向量空间，'
    '捕捉单词之间的语义关系。随后，循环神经网络(RNN)被广泛应用于序列建模任务，如机器翻译和文本生成。\n\n'
    '2017年，Vaswani等人[1]提出的Transformer架构通过自注意力机制解决了RNN难以并行化和捕捉长距离依赖的问题，成为NLP领域的里程碑。'
    '基于Transformer的预训练语言模型如BERT(Devlin et al., 2019)[2]进一步推动了NLP的发展，在多种任务上取得了突破性进展。'
)

# 添加主要应用部分
doc.add_heading('2. 深度学习在NLP中的主要应用', 1)

# 文本分类
doc.add_heading('2.1 文本分类', 2)
text_classification = doc.add_paragraph(
    '文本分类是NLP中的基础任务，包括情感分析、主题分类和垃圾邮件检测等。深度学习模型，特别是基于CNN的架构(Kim, 2014)[5]，'
    '在文本分类任务上取得了显著成功。预训练语言模型如BERT通过微调，进一步提高了文本分类的准确率(Devlin et al., 2019)[2]。'
)

# 机器翻译
doc.add_heading('2.2 机器翻译', 2)
mt = doc.add_paragraph(
    '神经机器翻译(NMT)是深度学习在NLP中最成功的应用之一。基于编码器-解码器架构的序列到序列模型和注意力机制显著提高了翻译质量。'
    'Transformer架构(Vaswani et al., 2017)[1]通过完全基于自注意力的设计，进一步推动了机器翻译的发展。'
)

# 添加预训练语言模型部分
doc.add_heading('3. 预训练语言模型的发展', 1)

# BERT及其变体
doc.add_heading('3.1 BERT及其变体', 2)
bert = doc.add_paragraph(
    'BERT(Bidirectional Encoder Representations from Transformers)是由Google AI团队开发的预训练语言模型(Devlin et al., 2019)[2]，它通过双向Transformer编码器学习上下文相关的词表示。'
    'BERT的预训练任务包括掩码语言模型(MLM)和下一句预测(NSP)，使其能够捕捉词级和句级的语义信息。'
)

# 添加结论部分
doc.add_heading('4. 结论', 1)
conclusion = doc.add_paragraph(
    '深度学习技术，特别是预训练语言模型，已经彻底改变了NLP领域的研究和应用。从词嵌入到Transformer架构，从RNN到BERT，NLP模型的能力和性能不断提升，在文本分类和机器翻译等任务上取得了显著进展。\n\n'
    '尽管取得了巨大成功，深度学习在NLP中仍面临计算资源和模型可解释性等挑战。未来的研究方向包括开发更高效的模型架构和提高模型透明度等。'
)

# 添加参考文献部分
doc.add_heading('参考文献', 1)
references = [
    '[1] Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, L., & Polosukhin, I. (2017). Attention is all you need. In Advances in Neural Information Processing Systems (pp. 5998-6008).',
    '[2] Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. In Proceedings of the 2019 Conference of the North American Chapter of the Association for Computational Linguistics: Human Language Technologies (pp. 4171-4186).',
    '[3] Goldberg, Y. (2017). Neural network methods for natural language processing. Synthesis Lectures on Human Language Technologies, 10(1), 1-309.',
    '[4] Mikolov, T., Sutskever, I., Chen, K., Corrado, G. S., & Dean, J. (2013). Distributed representations of words and phrases and their compositionality. In Advances in Neural Information Processing Systems (pp. 3111-3119).',
    '[5] Kim, Y. (2014). Convolutional neural networks for sentence classification. In Proceedings of the 2014 Conference on Empirical Methods in Natural Language Processing (pp. 1746-1751).'
]

for ref in references:
    doc.add_paragraph(ref)

# 保存文档
doc.save('d:\\Projects\\reference-agent\\mini_test_paper.docx')

print("迷你版Word文档已成功创建：d:\\Projects\\reference-agent\\mini_test_paper.docx")