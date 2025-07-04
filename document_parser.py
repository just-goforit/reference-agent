# document_parser.py
import re
import os
import logging
from typing import List, Dict, Tuple, Set, Optional

import docx
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from config import CITATION_PATTERNS
from utils import clean_text, logger, find_citation_context


class DocumentParser:
    """Word文档解析器，用于提取引用和参考文献"""
    
    def __init__(self, file_path: str):
        """初始化解析器
        
        Args:
            file_path: Word文档路径
        """
        self.file_path = file_path
        self.document = None
        self.text_content = ""
        self.content_without_references = ""
        self.paragraphs = []
        self.references_section = ""
        self.references = []
        self.citations = set()
        
        # 加载文档
        self._load_document()
    
    def _load_document(self) -> None:
        """加载Word文档"""
        try:
            if not os.path.exists(self.file_path):
                raise FileNotFoundError(f"文件不存在: {self.file_path}")
                
            self.document = docx.Document(self.file_path)
            logger.info(f"成功加载文档: {self.file_path}")
            
            # 提取文本内容
            self.paragraphs = self.document.paragraphs
            self.text_content = "\n".join([p.text for p in self.paragraphs])            

        except Exception as e:
            logger.error(f"加载文档失败: {str(e)}")
            raise
    
    def extract_references_section(self) -> str:
        """提取参考文献部分"""
        # 常见的参考文献部分标题
        ref_section_titles = [
            "References", "REFERENCES", "参考文献", "引用文献", 
            "Bibliography", "BIBLIOGRAPHY", "文献"
        ]
        
        # 查找参考文献部分
        found_references = False
        references_text = []
        
        for i, para in enumerate(self.paragraphs):
            text = para.text.strip()
            
            # 检查是否是参考文献部分的标题
            if not found_references:
                for title in ref_section_titles:
                    if text == title or text.startswith(title + ":"):
                        found_references = True
                        break
            
            # 如果已找到参考文献部分，收集文本
            elif found_references:
                # 检查是否到达下一个主要部分（通常是大写标题）
                if text and text.isupper() and len(text.split()) <= 3:
                    break
                
                # 添加非空段落
                if text:
                    references_text.append(text)
        
        self.references_section = "\n".join(references_text)
        logger.info(f"提取到参考文献部分，长度: {len(self.references_section)} 字符")

        # 排除参考文献部分的文本
        if self.references_section:
            # 找到参考文献部分的起始位置
            ref_start = self.text_content.find(self.references_section)
            if ref_start != -1:
                # 只使用参考文献部分之前的文本
                self.content_without_references = self.text_content[:ref_start]
        
        return self.references_section
    
    def parse_references(self) -> List[str]:
        """解析参考文献列表"""
        if not self.references_section:
            self.extract_references_section()
        
        # 如果仍然没有参考文献部分，尝试使用其他方法
        if not self.references_section:
            logger.warning("未找到明确的参考文献部分，尝试使用编号模式识别")
            self._parse_references_by_pattern()
            return self.references
        
        # 常见的参考文献格式模式
        reference_patterns = [
            # [1] Author, Title, Journal, Year
            r'\[\d+\]\s+.+',
            # 1. Author, Title, Journal, Year
            r'\d+\.\s+.+',
            # Author, Title, Journal, Year
            r'[A-Z][a-z]+,\s+[A-Z].+\(\d{4}\).+'
        ]
        
        # 尝试使用不同的模式分割参考文献
        for pattern in reference_patterns:
            references = re.findall(pattern, self.references_section)
            if references:
                self.references = [clean_text(ref) for ref in references]
                logger.info(f"使用模式 '{pattern}' 解析到 {len(self.references)} 条参考文献")
                return self.references
        
        # 如果没有匹配的模式，尝试按行分割
        lines = self.references_section.split('\n')
        potential_refs = []
        current_ref = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 检查是否是新的参考文献开始
            if re.match(r'^\[?\d+\]?\.?\s+', line) or re.match(r'^[A-Z][a-z]+,\s+[A-Z]', line):
                if current_ref:
                    potential_refs.append(clean_text(current_ref))
                current_ref = line
            else:
                current_ref += " " + line
        
        # 添加最后一条参考文献
        if current_ref:
            potential_refs.append(clean_text(current_ref))
        
        if potential_refs:
            self.references = potential_refs
            logger.info(f"通过行分析解析到 {len(self.references)} 条参考文献")
        else:
            logger.warning("无法解析参考文献，将整个部分作为一条参考文献")
            self.references = [clean_text(self.references_section)]
        
        return self.references
    
    def _parse_references_by_pattern(self) -> None:
        """通过模式识别参考文献"""
        # 查找文档末尾可能的参考文献
        last_paragraphs = self.paragraphs[-min(30, len(self.paragraphs)):]
        potential_refs = []
        
        for para in last_paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # 检查是否符合参考文献格式
            if (re.match(r'^\[?\d+\]?\.?\s+', text) or 
                re.match(r'^[A-Z][a-z]+,\s+[A-Z]', text) or
                "et al." in text or
                re.search(r'\(\d{4}\)', text)):
                potential_refs.append(clean_text(text))
        
        if potential_refs:
            self.references = potential_refs
            logger.info(f"通过模式识别解析到 {len(self.references)} 条参考文献")
    
    def extract_citations(self) -> Dict[str, List[str]]:
        """
        提取文档中的引用及其上下文
        
        Returns:
            引用字典，键为引用标识符，值为包含该引用的所有句子列表
        """
        citations = {}
        
        # 排除参考文献部分的文本
        text_without_references = self.text_content
        if self.references_section:
            # 找到参考文献部分的起始位置
            ref_start = self.text_content.find(self.references_section)
            if ref_start != -1:
                # 只使用参考文献部分之前的文本
                text_without_references = self.text_content[:ref_start]
        
        # 使用配置的引用模式
        for pattern in CITATION_PATTERNS:
            matches = re.findall(pattern, text_without_references)
            # matcher去重
            matches = list(set(matches))

            for citation in matches:
                # 过滤掉不应该被视为引用的文本
                # 排除 arXiv 标识符
                if citation.startswith('arXiv:'):
                    continue
                    
                # 排除年份格式 (2018) 或 (2019)
                if re.match(r'\(\d{4}\)', citation):
                    continue
                    
                # # 排除单独的数字后跟句点的情况，如 "1."
                # if re.match(r'^\d+\.?$', citation):
                #     continue
                
                logger.info(f"提取到引用: {citation}")

                # 使用 find_citation_context 函数获取包含引用的所有句子
                contexts = find_citation_context(text_without_references, citation)
                
                # 过滤上下文，只保留真正包含引用的完整句子
                valid_contexts = []
                for context in contexts:
                    # 确保上下文是完整句子，不只是引用标记本身
                    if len(context) > len(citation) and citation in context:
                        # 确保上下文不是参考文献部分的文本
                        if not any(ref in context for ref in self.references):
                            valid_contexts.append(context)
                
                if valid_contexts:
                    citations[citation] = valid_contexts
        
        # 更新实例的引用集合
        self.citations = set(citations.keys())
        logger.info(f"提取到 {len(citations)} 个引用")
        logger.info(f"提取到的引用: {citations}")
        
        return citations
    

    def get_document_metadata(self) -> Dict[str, str]:
        """获取文档元数据"""
        metadata = {}
        
        if self.document:
            core_properties = self.document.core_properties
            metadata["title"] = core_properties.title or "未知标题"
            metadata["author"] = core_properties.author or "未知作者"
            metadata["created"] = str(core_properties.created) if core_properties.created else "未知"
            metadata["modified"] = str(core_properties.modified) if core_properties.modified else "未知"
        
        return metadata